from docx import Document


def replace_placeholders_in_paragraph(paragraph, context):
    """Заменяет {KEY} на значение из context в параграфе"""
    for key, value in context.items():
        if value is None:
            value = ""
        placeholder = "{" + key + "}"
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, str(value))


def replace_placeholders_in_table(table, context):
    """Заменяет {KEY} на значение в таблице"""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_placeholders_in_paragraph(paragraph, context)


def fill_docx_template(template_path, context, output_docx_path):
    """Заполняет DOCX-шаблон данными и сохраняет новый файл"""
    doc = Document(template_path)

    #для параграфов
    for paragraph in doc.paragraphs:
        replace_placeholders_in_paragraph(paragraph, context)

    #ля таблиц
    for table in doc.tables:
        replace_placeholders_in_table(table, context)

    doc.save(output_docx_path)
    print("Заполненный DOCX сохранён!")


